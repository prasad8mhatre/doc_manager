from fastapi import APIRouter, UploadFile, Form, Depends
from sqlalchemy.orm import Session
from utils.database import get_db
from utils.embedding import generate_embedding, store_embeddings,retrieve_embeddings
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from models import QAModel
from PyPDF2 import PdfReader  # For PDF files
from docx import Document  # For DOCX files
from io import BytesIO
from fastapi import Body
import uuid
from langchain_core.documents import Document
router = APIRouter()


def extract_text_from_pdf(file: BytesIO) -> str:
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file: BytesIO) -> str:
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text
    return text

@router.post("/ingestion")
async def ingest_document(file: UploadFile, name: str = Form(...), db: Session = Depends(get_db)):
    
    if file.filename.endswith(".pdf"):
        content = extract_text_from_pdf(BytesIO(await file.read()))
    elif file.filename.endswith(".txt"):
        content = (await file.read()).decode("utf-8")
    elif file.filename.endswith(".docx"):
        content = extract_text_from_docx(BytesIO(await file.read()))
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF, TXT, and DOCX are supported.")
    
    user_id ='123'

    # Split the content into smaller parts, e.g., 1000 characters each
    chunk_size = 1000
    content_parts = [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]

    # For each part, generate an embedding and store it
    docs = []
    document_id = name+"-"+str(uuid.uuid4())  # or any unique ID generation logic you prefer
    stored_ids = []
    for part_id, part_content in enumerate(content_parts, 1):
        id = str(uuid.uuid4())
        stored_ids.append(id)
        docs.append(Document(page_content=part_content, metadata={"id":id, "userId": user_id, "document_id": document_id ,"document_name":name, "partId":part_id}))
    
    await store_embeddings(docs)
    
    return {"ids": stored_ids, "status": "Document ingested successfully in parts"}


@router.post("/retrieve")
async def retrieve_documents(body: dict = Body(...)):
    query = body.get("query")
    documents = retrieve_embeddings(query,2)
    if not documents:
        return {"status": "No matching documents found", "document_ids": []}
    return {"status": "Documents fetched", "documents": documents}

@router.post("/selection")
async def select_documents(document_ids: list, db: Session = Depends(get_db)):
    documents = retrieve_embeddings(db, document_ids)
    if not documents:
        return {"status": "No matching documents found", "document_ids": []}
    return {"status": "Documents selected", "document_ids": document_ids}
